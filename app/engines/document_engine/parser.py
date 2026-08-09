"""
SMART_AO V7 - parser.py
=======================
Moteur de parsing des documents DCE (PDF, ZIP, DOCX).
Extrait le texte et les métadonnées des pièces du marché.
"""
import logging
import os
from typing import Dict, Any, Optional, List
from datetime import datetime
import fitz  # PyMuPDF
import pdfplumber

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser multi-format pour documents de marchés publics."""

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt', '.zip']
    
    def __init__(self):
        self.parser_stats = {
            "documents_parsed": 0,
            "total_pages": 0,
            "errors": 0
        }

    async def parse_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Parse un fichier document et extrait son contenu.
        
        Args:
            file_path: Chemin complet vers le fichier
            
        Returns:
            Dict avec le contenu textuel et les métadonnées, ou None si échec
        """
        logger.info(f"Parsing du fichier : {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"Fichier non trouvé : {file_path}")
            return None

        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                result = await self._parse_pdf(file_path)
            elif file_ext == '.docx':
                result = await self._parse_docx(file_path)
            elif file_ext == '.txt':
                result = await self._parse_txt(file_path)
            elif file_ext == '.zip':
                result = await self._parse_zip(file_path)
            else:
                logger.warning(f"Format non supporté : {file_ext}")
                return None

            if result:
                self.parser_stats["documents_parsed"] += 1
                logger.debug(f"Parsing réussi : {file_path}")
            
            return result

        except Exception as e:
            logger.error(f"Erreur lors du parsing de {file_path}: {str(e)}")
            self.parser_stats["errors"] += 1
            return None

    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse un fichier PDF avec PyMuPDF et pdfplumber."""
        text_content = ""
        pages_count = 0
        metadata = {}

        try:
            # Extraction avec PyMuPDF (rapide)
            doc = fitz.open(file_path)
            pages_count = len(doc)
            
            for page_num in range(pages_count):
                page = doc[page_num]
                text_content += page.get_text()
            
            metadata = doc.metadata
            doc.close()
            
            # Extraction améliorée avec pdfplumber (tables)
            tables_data = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:5]:  # Premieres pages pour tables
                    tables = page.extract_tables()
                    if tables:
                        tables_data.extend(tables)

            return {
                "type": "pdf",
                "text": text_content,
                "pages": pages_count,
                "metadata": {
                    "title": metadata.get("title", ""),
                    "author": metadata.get("author", ""),
                    "creation_date": metadata.get("creationDate", ""),
                    "producer": metadata.get("producer", "")
                },
                "tables_count": len(tables_data),
                "parsed_at": datetime.utcnow().isoformat()
            }

        except Exception as e:
            logger.error(f"Erreur parsing PDF {file_path}: {str(e)}")
            raise

    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse un fichier DOCX."""
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_content = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                "type": "docx",
                "text": text_content,
                "pages": len(doc.paragraphs) // 10 + 1,  # Estimation
                "metadata": {},
                "parsed_at": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"Erreur parsing DOCX {file_path}: {str(e)}")
            raise

    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Parse un fichier texte simple."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            lines_count = len(text_content.splitlines())
            
            return {
                "type": "txt",
                "text": text_content,
                "pages": lines_count // 50 + 1,
                "metadata": {},
                "parsed_at": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"Erreur parsing TXT {file_path}: {str(e)}")
            raise

    async def _parse_zip(self, file_path: str) -> Dict[str, Any]:
        """Parse un fichier ZIP contenant plusieurs documents."""
        import zipfile
        
        try:
            zip_contents = []
            
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                
                for filename in file_list:
                    if any(filename.endswith(ext) for ext in self.SUPPORTED_FORMATS):
                        zip_contents.append({
                            "filename": filename,
                            "size": zip_ref.getinfo(filename).file_size
                        })

            return {
                "type": "zip",
                "text": "",
                "pages": len(zip_contents),
                "metadata": {"files": zip_contents},
                "parsed_at": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"Erreur parsing ZIP {file_path}: {str(e)}")
            raise

    def get_stats(self) -> Dict[str, int]:
        """Retourne les statistiques de parsing."""
        return self.parser_stats.copy()
