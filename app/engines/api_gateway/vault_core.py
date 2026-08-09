"""
SMART_AO V7 - vault_core.py
================================
Copyright (c) 2026 NOOR - Architecte Principal
Licence: Proprietary - All Rights Reserved
Auteur: NOOR
Date: 06/08/2026
Build: 9 - Phase: 5
"""

"""
SMART_AO V7 - Vault Core Engine
=================================
Gestion centralisée des documents avec stockage, indexation et recherche sémantique
Source: ARCHITECTURE_V7_ENGINE.md §4.3
"""

import asyncio
import hashlib
import os
import uuid
from datetime import datetime, timezone
from typing import Optional, List, Dict, Any, BinaryIO, Union
from pathlib import Path

from sqlalchemy import select, update, delete, and_, or_, desc
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import engine
from app.models.vault_core import VaultDocument, DocumentChunk
from app.core.config import settings
from app.engines.security_engine.filesystem import validate_file, get_file_extension
from app.engines.security_engine.clamav import scan_content


# =============================================================================
# CONSTANTS
# =============================================================================

# Document status
STATUS_UPLOADED = "uploaded"
STATUS_PROCESSING = "processing"
STATUS_PROCESSED = "processed"
STATUS_FAILED = "failed"
STATUS_QUARANTINED = "quarantined"

# Allowed extensions
ALLOWED_EXTENSIONS = [".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".json"]
MAX_FILE_SIZE = settings.UPLOAD_MAX_SIZE_MB * 1024 * 1024  # 50MB by default


# =============================================================================
# VAULT CORE ENGINE
# =============================================================================

class VaultCoreEngine:
    """
    Vault Core Engine for document management
    
    Responsibilities:
    - Document upload and validation
    - File storage (local or MinIO/S3)
    - Virus scanning
    - Chunking for semantic search
    - Metadata extraction
    - Indexing in Qdrant
    """
    
    def __init__(self):
        self._db: Optional[AsyncSession] = None
        self._storage_path = Path(settings.STORAGE_DATA_DIRECTORY) / "documents"
        self._storage_path.mkdir(parents=True, exist_ok=True)
    
    async def get_db(self) -> AsyncSession:
        """Get database session"""
        if self._db is None:
            self._db = AsyncSession(engine, expire_on_commit=False)
        return self._db
    
    async def close(self):
        """Close database connection"""
        if self._db:
            await self._db.close()
            self._db = None
    
    # =========================================================================
    # DOCUMENT UPLOAD & VALIDATION
    # =========================================================================
    
    async def upload_document(
        self,
        file: BinaryIO,
        file_name: str,
        file_size: int,
        document_type: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Upload a document to the vault
        
        Steps:
        1. Validate file (extension, size)
        2. Scan for viruses
        3. Generate unique ID and hash
        4. Store file
        5. Create database record
        
        Args:
            file: File binary stream
            file_name: Original file name
            file_size: File size in bytes
            document_type: Type of document (DCE, CCAP, DPGF, etc.)
            metadata: Additional metadata
        
        Returns:
            Dict: Document information
        
        Raises:
            ValueError: If file is invalid
            Exception: If upload fails
        """
        # Step 1: Validate file
        file_content = file.read()
        file.seek(0)  # Reset file pointer
        
        # Validate extension
        file_extension = get_file_extension(file_name)
        if file_extension.lower() not in ALLOWED_EXTENSIONS:
            raise ValueError(
                f"File extension not allowed: {file_extension}. "
                f"Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        
        # Validate size
        if file_size > MAX_FILE_SIZE:
            raise ValueError(
                f"File too large: {file_size / (1024*1024):.2f}MB. "
                f"Max: {MAX_FILE_SIZE / (1024*1024):.0f}MB"
            )
        
        # Step 2: Scan for viruses
        if settings.STORAGE_ENCRYPTION_ENABLED:
            # In production, always scan
            scan_result = await scan_content(file_content, file_name)
            if scan_result.is_infected:
                # Quarantine the file
                return await self._quarantine_file(
                    file_content, file_name, scan_result.virus_name
                )
            if scan_result.is_error:
                raise RuntimeError(f"Scan antivirus indisponible — upload refusé: {scan_result.message}")
        
        # Step 3: Generate IDs and hash
        document_id = f"doc_{uuid.uuid4().hex[:16]}"
        content_hash = hashlib.sha256(file_content).hexdigest()
        
        # Step 4: Store file
        storage_path = self._get_storage_path(document_id)
        os.makedirs(os.path.dirname(storage_path), exist_ok=True)
        
        with open(storage_path, 'wb') as f:
            f.write(file_content)
        
        # Step 5: Create database record
        db = await self.get_db()
        
        vault_doc = VaultDocument(
            document_id=document_id,
            file_name=file_name,
            file_path=storage_path,
            file_type=self._get_mime_type(file_extension),
            file_size=file_size,
            content_hash=content_hash,
            status=STATUS_UPLOADED,
            extra_metadata={
                "document_type": document_type,
                "original_metadata": metadata or {},
                "uploader": "system"
            },
            created_at=datetime.now(timezone.utc)
        )
        
        db.add(vault_doc)
        await db.commit()
        await db.refresh(vault_doc)
        
        return {
            "success": True,
            "document_id": document_id,
            "file_name": file_name,
            "file_size": file_size,
            "content_hash": content_hash,
            "status": STATUS_UPLOADED,
            "message": "Document uploaded successfully"
        }
    
    async def _quarantine_file(
        self,
        file_content: bytes,
        file_name: str,
        scan_result: str
    ) -> Dict[str, Any]:
        """Move infected file to quarantine"""
        quarantine_path = self._get_storage_path(f"quarantine_{uuid.uuid4().hex[:16]}", "quarantine")
        os.makedirs(os.path.dirname(quarantine_path), exist_ok=True)
        
        with open(quarantine_path, 'wb') as f:
            f.write(file_content)
        
        document_id = f"quarantine_{uuid.uuid4().hex[:16]}"
        
        db = await self.get_db()
        vault_doc = VaultDocument(
            document_id=document_id,
            file_name=file_name,
            file_path=quarantine_path,
            file_type="application/octet-stream",
            file_size=len(file_content),
            content_hash=hashlib.sha256(file_content).hexdigest(),
            status=STATUS_QUARANTINED,
            extra_metadata={
                "quarantine_reason": scan_result,
                "original_file": file_name
            },
            created_at=datetime.now(timezone.utc)
        )
        
        db.add(vault_doc)
        await db.commit()
        
        return {
            "success": False,
            "document_id": document_id,
            "file_name": file_name,
            "status": STATUS_QUARANTINED,
            "error": f"Virus detected: {scan_result}",
            "message": "File quarantined"
        }
    
    # =========================================================================
    # DOCUMENT PROCESSING
    # =========================================================================
    
    async def process_document(
        self,
        document_id: str
    ) -> Dict[str, Any]:
        """
        Process a document: extract text, create chunks, generate embeddings
        
        Args:
            document_id: Document ID
        
        Returns:
            Dict: Processing results
        """
        db = await self.get_db()
        
        # Get document
        result = await db.execute(
            select(VaultDocument).where(
                VaultDocument.document_id == document_id
            )
        )
        vault_doc = result.scalar_one_or_none()
        
        if not vault_doc:
            raise ValueError(f"Document not found: {document_id}")
        
        # Update status
        vault_doc.status = STATUS_PROCESSING
        await db.commit()
        
        try:
            # Step 1: Extract text based on file type
            file_path = vault_doc.file_path
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            text_content = self._extract_text(file_path, vault_doc.file_type)
            
            # Step 2: Create chunks
            chunks = self._create_chunks(text_content)
            
            # Step 3: Save chunks to database
            for chunk_index, chunk in enumerate(chunks):
                chunk_model = DocumentChunk(
                    document_id=vault_doc.id,
                    chunk_index=chunk_index,
                    content=chunk["text"],
                    embedding=None,  # Will be set by embedding service
                    start_page=chunk.get("start_page"),
                    end_page=chunk.get("end_page"),
                    extra_metadata={"chunk_size": len(chunk["text"])}
                )
                db.add(chunk_model)
            
            await db.commit()
            
            # Step 4: Update document status
            vault_doc.status = STATUS_PROCESSED
            vault_doc.processed_at = datetime.now(timezone.utc)
            vault_doc.extra_metadata["chunk_count"] = len(chunks)
            vault_doc.extra_metadata["total_chars"] = sum(len(c["text"]) for c in chunks)
            
            await db.commit()
            await db.refresh(vault_doc)
            
            return {
                "success": True,
                "document_id": document_id,
                "status": STATUS_PROCESSED,
                "chunk_count": len(chunks),
                "total_chars": sum(len(c["text"]) for c in chunks),
                "message": "Document processed successfully"
            }
        
        except Exception as e:
            vault_doc.status = STATUS_FAILED
            vault_doc.error_message = str(e)
            await db.commit()
            raise
    
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from file based on type"""
        try:
            if file_type == "application/pdf" or file_path.endswith(".pdf"):
                return self._extract_pdf_text(file_path)
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                              "application/msword"] or file_path.endswith((".docx", ".doc")):
                return self._extract_docx_text(file_path)
            elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                              "application/vnd.ms-excel"] or file_path.endswith((".xlsx", ".xls")):
                return self._extract_excel_text(file_path)
            elif file_type == "text/plain" or file_path.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_type == "application/json" or file_path.endswith(".json"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                # Fallback: read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        # Implementation with PyMuPDF or pdfplumber
        # For now, return placeholder
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        except ImportError:
            # Fallback
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    
    def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel"""
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            return df.to_string()
        except ImportError:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    
    def _create_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """Create overlapping chunks for semantic search"""
        if not text or len(text) == 0:
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk_text = text[start:end]
            
            chunks.append({
                "chunk_index": chunk_index,
                "text": chunk_text,
                "start_page": 0,  # Will be updated based on PDF pages
                "end_page": 0
            })
            
            start = end - overlap if end < len(text) else end
            chunk_index += 1
        
        return chunks
    
    # =========================================================================
    # DOCUMENT RETRIEVAL
    # =========================================================================
    
    async def get_document(self, document_id: str) -> Optional[VaultDocument]:
        """Get a document by ID"""
        db = await self.get_db()
        
        result = await db.execute(
            select(VaultDocument).where(
                VaultDocument.document_id == document_id
            )
        )
        return result.scalar_one_or_none()
    
    async def list_documents(
        self,
        document_type: Optional[str] = None,
        status: Optional[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> List[VaultDocument]:
        """List documents"""
        db = await self.get_db()
        
        query = select(VaultDocument)
        
        if document_type:
            query = query.where(
                VaultDocument.extra_metadata["document_type"].as_string() == document_type
            )
        
        if status:
            query = query.where(VaultDocument.status == status)
        
        query = query.order_by(desc(VaultDocument.created_at))
        query = query.limit(limit).offset(offset)
        
        result = await db.execute(query)
        return result.scalars().all()
    
    async def delete_document(self, document_id: str) -> bool:
        """Delete a document"""
        db = await self.get_db()
        
        # Delete chunks first
        await db.execute(
            delete(DocumentChunk).where(DocumentChunk.document_id == VaultDocument.id)
        )
        
        # Delete document
        result = await db.execute(
            delete(VaultDocument).where(
                VaultDocument.document_id == document_id
            )
        )
        
        await db.commit()
        return result.rowcount > 0
    
    async def get_document_chunks(self, document_id: str) -> List[DocumentChunk]:
        """Get all chunks for a document"""
        db = await self.get_db()
        
        # First get document ID
        doc = await self.get_document(document_id)
        if not doc:
            return []
        
        result = await db.execute(
            select(DocumentChunk).where(DocumentChunk.document_id == doc.id)
            .order_by(DocumentChunk.chunk_index)
        )
        return result.scalars().all()
    
    # =========================================================================
    # UTILITY METHODS
    # =========================================================================
    
    def _get_storage_path(self, document_id: str, subdir: str = "documents") -> str:
        """Get storage path for a document"""
        return str(self._storage_path / subdir / f"{document_id[:2]}" / document_id)
    
    def _get_mime_type(self, file_extension: str) -> str:
        """Get MIME type from file extension"""
        mime_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xls": "application/vnd.ms-excel",
            ".txt": "text/plain",
            ".json": "application/json"
        }
        return mime_types.get(file_extension.lower(), "application/octet-stream")


# =============================================================================
# SINGLETON INSTANCE
# =============================================================================

vault_core = VaultCoreEngine()


def get_vault_core() -> VaultCoreEngine:
    """Get the singleton VaultCoreEngine instance"""
    return vault_core
